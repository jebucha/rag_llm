import os
import re
import traceback
import time
from typing import List, Tuple, Dict, Any, Optional
from concurrent.futures import ProcessPoolExecutor, as_completed
import multiprocessing
import tempfile
import queue
import threading
import concurrent.futures
import streamlit as st
import fitz  # PyMuPDF
import docx
import chromadb
from ollama import Client as OllamaClient
from sentence_transformers import SentenceTransformer
from transformers import AutoTokenizer
import numpy as np
from tqdm.auto import tqdm
from threading import Thread

# This must be the first Streamlit command
st.set_page_config(
    page_title="RAG Document Assistant",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ----------------------------
# Caching Resources
# ----------------------------

@st.cache_resource
def load_embedding_model(model_name: str = "all-MiniLM-L6-v2") -> SentenceTransformer:
    """
    Load and cache the SentenceTransformer model.
    """
    return SentenceTransformer(model_name)

@st.cache_resource
def load_tokenizer(model_choice: str) -> AutoTokenizer:
    """
    Load and cache the tokenizer based on the given model choice.
    """
    tokenizer_map = {
        "mistral": "intfloat/e5-mistral-7b-instruct",
        "gemma3": "google/generative-ai",
        "cogito": "NousResearch/Llama-2-7b-hf",
        "ALIENTELLIGENCE/contractanalyzerv2": "sentence-transformers/all-MiniLM-L6-v2",
    }
    model_id = tokenizer_map.get(model_choice, "sentence-transformers/all-MiniLM-L6-v2")
    return AutoTokenizer.from_pretrained(model_id)

# Load the embedding model once for the session
model = load_embedding_model()

# Setup persistent Chroma client and Ollama client
persist_path = "./chroma_db"
client = chromadb.PersistentClient(path=persist_path)
ollama = OllamaClient()

# ----------------------------
# Helper Functions for Parallel Processing
# ----------------------------

def extract_text_from_pdf_file(file_path: str) -> Tuple[str, bool]:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in doc])
        return text, True
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}", False

def extract_text_from_txt_file(file_path: str) -> Tuple[str, bool]:
    """Extract text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read(), True
    except Exception as e:
        return f"Error extracting text from TXT: {str(e)}", False

def extract_text_from_docx_file(file_path: str) -> Tuple[str, bool]:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs]), True
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}", False

def smart_chunk(text: str, min_length: int = 40) -> List[str]:
    """
    Chunk text using a simple heuristic that splits on two or more newlines
    and filters out chunks that are too short.
    """
    chunks = [chunk.strip() for chunk in re.split(r'\n{2,}', text) if len(chunk.strip()) > min_length]
    return chunks

def smart_chunk_markdown(text: str) -> List[str]:
    """
    Chunk markdown text by splitting on patterns where list items or headers occur.
    """
    return [chunk.strip() for chunk in re.split(r'\n(?=(\d+\.\s|\-|\#))', text) if chunk.strip()]

def process_file_for_parallel(args):
    """
    Process a single file in a separate process.

    Args:
        args: Tuple containing (file_path, file_type, file_name)

    Returns:
        Dict with processing results including embeddings
    """
    file_path, file_type, file_name = args

    result = {
        "filename": file_name,
        "success": False,
        "chunks": [],
        "embeddings": [],
        "ids": [],
        "metadatas": [],
        "error": None
    }

    try:
        # Extract text based on file type
        if file_type == "pdf":
            text, success = extract_text_from_pdf_file(file_path)
        elif file_type == "txt":
            text, success = extract_text_from_txt_file(file_path)
        elif file_type == "docx":
            text, success = extract_text_from_docx_file(file_path)
        elif file_type == "markdown":
            text, success = extract_text_from_txt_file(file_path)
        else:
            return {**result, "error": f"Unsupported file type: {file_type}"}

        if not success:
            return {**result, "error": text}  # text contains error message

        # Split the text into semantic chunks
        chunks = smart_chunk_markdown(text) if file_type == "markdown" else smart_chunk(text)

        if not chunks:
            return {**result, "error": "No valid chunks extracted"}

        # Generate embeddings
        # Load model inside the process to avoid sharing issues
        local_model = SentenceTransformer("all-MiniLM-L6-v2")
        embeddings = local_model.encode(chunks)

        base = os.path.splitext(file_name)[0]
        ids = [f"{base}_{file_type}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [{"source": file_name, "type": file_type} for _ in chunks]

        return {
            **result,
            "success": True,
            "chunks": chunks,
            "embeddings": embeddings.tolist(),
            "ids": ids,
            "metadatas": metadatas
        }

    except Exception as e:
        error_details = traceback.format_exc()
        return {**result, "error": f"{str(e)}\n{error_details}"}

def batch_add_to_collection(collection, results_queue, batch_size=1000):
    """
    Worker function to add processed results to the collection in batches.

    Args:
        collection: ChromaDB collection
        results_queue: Queue containing processing results
        batch_size: Number of documents to add in a single batch
    """
    all_chunks = []
    all_embeddings = []
    all_ids = []
    all_metadatas = []

    while True:
        try:
            result = results_queue.get(timeout=1)

            if result == "DONE":
                break

            if result["success"]:
                all_chunks.extend(result["chunks"])
                all_embeddings.extend(result["embeddings"])
                all_ids.extend(result["ids"])
                all_metadatas.extend(result["metadatas"])

                # If we've reached the batch size, add to collection
                if len(all_chunks) >= batch_size:
                    collection.add(
                        documents=all_chunks,
                        embeddings=all_embeddings,
                        ids=all_ids,
                        metadatas=all_metadatas
                    )
                    all_chunks = []
                    all_embeddings = []
                    all_ids = []
                    all_metadatas = []

            results_queue.task_done()

        except queue.Empty:
            # If queue is empty for a while, check if there's anything to add
            if all_chunks:
                collection.add(
                    documents=all_chunks,
                    embeddings=all_embeddings,
                    ids=all_ids,
                    metadatas=all_metadatas
                )
                all_chunks = []
                all_embeddings = []
                all_ids = []
                all_metadatas = []

    # Add any remaining documents
    if all_chunks:
        collection.add(
            documents=all_chunks,
            embeddings=all_embeddings,
            ids=all_ids,
            metadatas=all_metadatas
        )

# ----------------------------
# Original Helper Functions (for Streamlit file objects)
# ----------------------------

def extract_text_from_pdf(file) -> Tuple[str, bool]:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = "\n".join([page.get_text() for page in doc])
        return text, True
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}", False

def extract_text_from_txt(file) -> Tuple[str, bool]:
    """Extract text from a TXT file."""
    try:
        return file.read().decode("utf-8"), True
    except Exception as e:
        return f"Error extracting text from TXT: {str(e)}", False

def extract_text_from_docx(file) -> Tuple[str, bool]:
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs]), True
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}", False

def process_file(file, collection, file_type: str) -> Dict[str, Any]:
    """
    Process a single file and add it to the collection.
    Returns a dictionary with processing results.
    """
    result = {
        "filename": file.name,
        "success": False,
        "chunks": 0,
        "error": None
    }

    try:
        # Extract text based on file type
        if file_type == "pdf":
            text, success = extract_text_from_pdf(file)
        elif file_type == "txt":
            text, success = extract_text_from_txt(file)
        elif file_type == "docx":
            text, success = extract_text_from_docx(file)
        elif file_type == "markdown":
            text, success = extract_text_from_txt(file)
        else:
            return {**result, "error": f"Unsupported file type: {file_type}"}

        if not success:
            return {**result, "error": text}  # text contains error message

        # Split the text into semantic chunks
        chunks = smart_chunk_markdown(text) if file_type == "markdown" else smart_chunk(text)

        if not chunks:
            return {**result, "error": "No valid chunks extracted"}

        # Generate embeddings
        embeddings = model.encode(chunks)
        base = os.path.splitext(file.name)[0]
        ids = [f"{base}_{file_type}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [{"source": file.name, "type": file_type} for _ in chunks]

        # Add to collection
        collection.add(documents=chunks, embeddings=embeddings.tolist(), ids=ids, metadatas=metadatas)

        return {
            **result,
            "success": True,
            "chunks": len(chunks)
        }

    except Exception as e:
        error_details = traceback.format_exc()
        return {**result, "error": f"{str(e)}\n{error_details}"}

# ----------------------------
# Application Tabs
# ----------------------------

tab_query, tab_ingest_pdf, tab_ingest_folder, tab_ingest_md, tab_verify, tab_errors = st.tabs([
    "Ask Questions",
    "Ingest Files",
    "Ingest Folder",
    "Ingest Markdown",
    "Verify Documents",
    "Errors"
])

# Store errors in session state for persistence
if 'ingestion_errors' not in st.session_state:
    st.session_state.ingestion_errors = []

# ----------------------------
# Tab 1: Ask Questions
# ----------------------------

with tab_query:
    st.header("Ask Questions About Your Documents")
    question = st.text_input("Enter your question:")
    top_n = st.slider("How many chunks to retrieve?", 1, 10, 3)
    model_choice = st.selectbox("Choose LLM model", ["mistral", "cogito", "gemma3", "ALIENTELLIGENCE/contractanalyzerv2"])

    if question:
        collection = client.get_or_create_collection("tutorial_docs")
        # Generate embedding for the query
        embedding = model.encode(question)
        results = collection.query(query_embeddings=[embedding.tolist()], n_results=top_n)
        docs = results.get("documents", [[]])[0]

        if not docs:
            st.warning("No matching documents found.")
        else:
            # Build the prompt from retrieved document chunks
            ranked_chunks = "\n\n".join([f"[Chunk {i+1}]\n{doc}" for i, doc in enumerate(docs)])
            prompt = (
                "You are a highly helpful assistant.\n\n"
                "Use the following retrieved text chunks to answer the question.\n\n"
                "Chunks:\n\"\"\"\n" +
                ranked_chunks +
                "\n\"\"\"\n\n" +
                f"Question: {question}\n" +
                "Answer:"
            )

            # Load and cache the tokenizer for the chosen model
            try:
                tokenizer = load_tokenizer(model_choice)
                token_count = len(tokenizer.encode(prompt))
            except Exception as e:
                st.warning(f"Tokenizer error: {e}")
                token_count = 0

            # Generate the response using the LLM via Ollama
            response = ollama.chat(
                model=model_choice,
                messages=[{"role": "user", "content": prompt}]
            )
            answer = response["message"]["content"].strip()

            # Display the answer and additional information
            st.subheader("Answer")
            st.write(answer)
            st.markdown(f"Prompt token count: {token_count}")

            with st.expander("Show retrieved context"):
                for i, doc in enumerate(docs):
                    st.markdown(f"Chunk {i+1}")
                    st.code(doc)

            with st.expander("Show prompt sent to model"):
                st.code(prompt)

# ----------------------------
# Tab 2: Ingest PDF/TXT/DOCX Files
# ----------------------------

with tab_ingest_pdf:
    st.header("Ingest PDF, TXT, or DOCX Documents")
    uploaded_files = st.file_uploader("Upload files", type=["pdf", "txt", "docx"], accept_multiple_files=True)

    if uploaded_files:
        collection = client.get_or_create_collection("tutorial_docs")

        # Create a progress bar for overall progress
        progress_bar = st.progress(0.0)

        # Create a status area for current file
        current_file_text = st.empty()

        # Process files with progress tracking
        total_files = len(uploaded_files)
        successful_files = 0
        failed_files = 0
        total_chunks = 0

        for i, file in enumerate(uploaded_files):
            # Update current file display
            current_file_text.text(f"Ingesting: {file.name} ({i+1}/{total_files})")

            # Update progress bar
            progress_bar.progress((i) / total_files)

            # Determine file type
            ext = os.path.splitext(file.name)[1].lower()
            if ext == ".pdf":
                file_type = "pdf"
            elif ext == ".txt":
                file_type = "txt"
            elif ext == ".docx":
                file_type = "docx"
            else:
                failed_files += 1
                st.session_state.ingestion_errors.append({
                    "filename": file.name,
                    "error": f"Unsupported file type: {ext}"
                })
                continue

            # Process the file
            result = process_file(file, collection, file_type)

            if result["success"]:
                successful_files += 1
                total_chunks += result["chunks"]
            else:
                failed_files += 1
                st.session_state.ingestion_errors.append({
                    "filename": file.name,
                    "error": result["error"]
                })

            # Update progress bar
            progress_bar.progress((i+1) / total_files)

        # Clear the current file text
        current_file_text.empty()

        # Complete the progress bar
        progress_bar.progress(1.0)

        # Show final summary
        st.info(f"Ingestion complete. Successfully processed {successful_files} of {total_files} files ({total_chunks} total chunks).")

        if failed_files > 0:
            st.warning(f"{failed_files} files failed to process. See the 'Errors' tab for details.")

# ----------------------------
# Tab 3: Ingest Folder (NEW - High-Speed Parallel Processing)
# ----------------------------

# Function to monitor system resources
def monitor_resources(resource_status, interval=2.0):
    """Monitor system resources and update status"""
    try:
        import psutil
        while not st.session_state.stop_ingestion and st.session_state.ingestion_running:
            try:
                # Get current memory and CPU usage
                memory_percent = psutil.virtual_memory().percent
                cpu_percent = psutil.cpu_percent(interval=0.5)

                # Update resource status
                resource_status["memory"] = memory_percent
                resource_status["cpu"] = cpu_percent

                # Determine if we need to throttle
                if memory_percent > 75:  # Default to 75% memory threshold
                    resource_status["throttle"] = True
                    resource_status["message"] = f"Memory usage high ({memory_percent}%)"
                elif cpu_percent > 85:  # Default to 85% CPU threshold
                    resource_status["throttle"] = True
                    resource_status["message"] = f"CPU usage high ({cpu_percent}%)"
                else:
                    resource_status["throttle"] = False
                    resource_status["message"] = "Resources OK"

                # Sleep for a bit
                time.sleep(interval)
            except Exception as e:
                # Don't crash the monitor thread
                resource_status["message"] = f"Error monitoring: {str(e)}"
                time.sleep(interval)
    except ImportError:
        # psutil not available
        resource_status["message"] = "psutil not available"
        while not st.session_state.stop_ingestion and st.session_state.ingestion_running:
            time.sleep(5)  # Just sleep and do nothing

# Function to manage the worker pool size
def adjust_workers(executor, resource_status, initial_workers):
    """Dynamically adjust the number of workers based on system load"""
    current_workers = initial_workers
    min_workers = 1

    while not st.session_state.stop_ingestion and st.session_state.ingestion_running:
        try:
            # If throttling is needed, reduce workers
            if resource_status.get("throttle", False):
                if current_workers > min_workers:
                    current_workers = max(min_workers, current_workers - 1)
                    executor._max_workers = current_workers
                    st.session_state.active_workers = current_workers
            # If resources are OK and we're below initial workers, increase
            elif current_workers < initial_workers:
                current_workers += 1
                executor._max_workers = current_workers
                st.session_state.active_workers = current_workers

            # Sleep for a bit
            time.sleep(5)
        except Exception:
            # Don't crash the adjuster thread
            time.sleep(5)

def stop_ingestion():
    st.session_state.stop_ingestion = True
    st.warning("Stopping ingestion process... This may take a moment.")

    # Force stop after 10 seconds if it doesn't stop naturally
    def force_stop():
        time.sleep(10)
        if st.session_state.ingestion_running:
            st.session_state.ingestion_running = False
            st.error("Forced stop after timeout. The application may need to be refreshed.")

    force_stop_thread = threading.Thread(target=force_stop)
    force_stop_thread.daemon = True
    force_stop_thread.start()

with tab_ingest_folder:
    st.header("High-Speed Folder Ingestion")

    col1, col2 = st.columns(2)

    with col1:
        folder_path = st.text_input("Enter folder path containing documents:")

        # Show the validate button
        validate_button = st.button("Validate Folder Path")

        # Only validate if the button is clicked and there's a path
        if validate_button:
            if not folder_path:
                st.error("Please enter a folder path first")
            elif not os.path.exists(folder_path):
                st.error(f"Path does not exist: {folder_path}")
            elif not os.path.isdir(folder_path):
                st.error(f"Path exists but is not a directory: {folder_path}")
            elif not os.access(folder_path, os.R_OK):
                st.error(f"Permission denied: Cannot read from {folder_path}")
            else:
                # Count files
                file_count = 0
                matching_files = 0
                file_types_found = set()

                with st.spinner("Scanning folder..."):
                    for root, _, files in os.walk(folder_path):
                        for file in files:
                            file_count += 1
                            ext = os.path.splitext(file)[1].lower()[1:]
                            if ext in ["pdf", "txt", "docx", "md"]:
                                matching_files += 1
                                file_types_found.add(ext)

                st.success(f"✅ Valid folder path! Found {file_count} total files, including {matching_files} processable files.")
                if file_types_found:
                    st.info(f"File types found: {', '.join(file_types_found)}")

    with col2:
        # Default to a more conservative number of workers
        default_workers = max(1, min(multiprocessing.cpu_count() - 1, 4))
        max_workers = st.slider("Number of parallel workers", 
                               min_value=1, 
                               max_value=multiprocessing.cpu_count(), 
                               value=default_workers)

        batch_size = st.slider("Batch size for adding to database", 
                              min_value=100, 
                              max_value=5000, 
                              value=1000,
                              step=100)

    file_types = st.multiselect("File types to process", 
                               ["pdf", "txt", "docx", "md"], 
                               default=["pdf", "txt", "docx", "md"])

    # Add example path help
    st.caption("Example path format: /Users/username/Documents/my_files or C:\\Users\\username\\Documents\\my_files")

    # Create a placeholder for the stop button
    stop_button_placeholder = st.empty()

    # Store ingestion state in session state
    if 'ingestion_running' not in st.session_state:
        st.session_state.ingestion_running = False

    # Create a flag for stopping the process
    if 'stop_ingestion' not in st.session_state:
        st.session_state.stop_ingestion = False

    # Create a variable for current active workers
    if 'active_workers' not in st.session_state:
        st.session_state.active_workers = max_workers

    def stop_ingestion():
        st.session_state.stop_ingestion = True
        st.warning("Stopping ingestion process... This may take a moment.")

        # Force stop after 10 seconds if it doesn't stop naturally
        def force_stop():
            time.sleep(10)
            if st.session_state.ingestion_running:
                st.session_state.ingestion_running = False
                st.error("Forced stop after timeout. The application may need to be refreshed.")

        force_stop_thread = Thread(target=force_stop)
        force_stop_thread.daemon = True
        force_stop_thread.start()

    if st.button("Start High-Speed Ingestion"):
        # Reset stop flag
        st.session_state.stop_ingestion = False
        st.session_state.ingestion_running = True
        st.session_state.active_workers = max_workers

        if not folder_path:
            st.error("Please enter a folder path")
            st.session_state.ingestion_running = False
        elif not os.path.exists(folder_path):
            st.error(f"Path does not exist: {folder_path}")
            st.session_state.ingestion_running = False
        elif not os.path.isdir(folder_path):
            st.error(f"Path exists but is not a directory: {folder_path}")
            st.session_state.ingestion_running = False
        elif not os.access(folder_path, os.R_OK):
            st.error(f"Permission denied: Cannot read from {folder_path}")
            st.session_state.ingestion_running = False
        else:
            # Get all files in the directory with the specified extensions
            all_files = []
            try:
                # Display a spinner while scanning for files
                with st.spinner(f"Scanning {folder_path} for files..."):
                    file_count = 0
                    for root, _, files in os.walk(folder_path):
                        for file in files:
                            file_count += 1
                            ext = os.path.splitext(file)[1].lower()[1:]  # Remove the dot
                            if ext in file_types:
                                file_path = os.path.join(root, file)
                                file_type = "markdown" if ext == "md" else ext
                                all_files.append((file_path, file_type, file))

                st.info(f"Scanned {file_count} total files, found {len(all_files)} matching files with types: {', '.join(file_types)}")

                if not all_files:
                    st.warning(f"No matching files found in {folder_path}")
                    st.session_state.ingestion_running = False
                else:
                    total_files = len(all_files)
                    st.info(f"Found {total_files} files to process. Starting high-speed ingestion...")

                    # Show stop button
                    stop_button_placeholder.button("Stop Ingestion", on_click=stop_ingestion, key="stop_ingestion_button")

                    # Create a progress bar and status text
                    progress_bar = st.progress(0.0)
                    status_text = st.empty()
                    metrics_container = st.container()

                    # Initialize metrics display
                    with metrics_container:
                        cols = st.columns(4)
                        processed_metric = cols[0].empty()
                        successful_metric = cols[1].empty()
                        failed_metric = cols[2].empty()
                        chunks_metric = cols[3].empty()

                        processed_metric.metric("Processed", "0")
                        successful_metric.metric("Successful", "0")
                        failed_metric.metric("Failed", "0")
                        chunks_metric.metric("Total Chunks", "0")

                    # Create a collection
                    collection = client.get_or_create_collection("tutorial_docs")

                    # Create a queue for results
                    results_queue = queue.Queue()

                    # Start a worker thread to add results to the collection
                    add_thread = threading.Thread(
                        target=batch_add_to_collection,
                        args=(collection, results_queue, batch_size)
                    )
                    add_thread.daemon = True
                    add_thread.start()

                    # Create a dictionary to track resource status
                    resource_status = {
                        "memory": 0,
                        "cpu": 0,
                        "throttle": False,
                        "message": "Starting..."
                    }

                    # Start resource monitoring thread
                    monitor_thread = Thread(target=monitor_resources, args=(resource_status,))
                    monitor_thread.daemon = True
                    monitor_thread.start()

                    # Process files in parallel
                    start_time = time.time()
                    processed_count = 0
                    successful_count = 0
                    failed_count = 0
                    total_chunks = 0

                    # Create a list to track errors during processing
                    processing_errors = []

                    # Track throttling pauses
                    throttle_count = 0
                    total_throttle_time = 0

                    # Simple throttling - pause occasionally
                    last_pause_time = time.time()
                    pause_interval = 50  # Pause after every 50 files

                    with ProcessPoolExecutor(max_workers=max_workers) as executor:
                        # Start worker adjuster thread
                        adjuster_thread = Thread(
                            target=adjust_workers, 
                            args=(executor, resource_status, max_workers)
                        )
                        adjuster_thread.daemon = True
                        adjuster_thread.start()

                        # Submit initial batch of tasks (limit to prevent overload)
                        initial_batch_size = min(100, len(all_files))
                        futures = {}
                        for i in range(initial_batch_size):
                            if i < len(all_files):
                                future = executor.submit(process_file_for_parallel, all_files[i])
                                futures[future] = all_files[i]

                        next_file_index = initial_batch_size

                        # Process results as they complete and submit new tasks
                        while futures and not st.session_state.stop_ingestion:
                            # Check if we need to pause due to high resource usage
                            should_throttle = False

                            # Simple throttling - pause every N files
                            current_time = time.time()
                            if processed_count > 0 and processed_count % pause_interval == 0 and current_time - last_pause_time > 60:
                                should_throttle = True
                                throttle_message = "Regular pause to prevent overload"
                                last_pause_time = current_time

                            if should_throttle:
                                throttle_start = time.time()
                                status_text.warning(f"⚠️ Throttling: {throttle_message} - Pausing for 5 seconds")
                                time.sleep(5)  # Pause for 5 seconds
                                throttle_count += 1
                                total_throttle_time += time.time() - throttle_start
                                continue

                            # Wait for the next future to complete
                            done, _ = concurrent.futures.wait(
                                futures, 
                                return_when=concurrent.futures.FIRST_COMPLETED,
                                timeout=0.5  # Add timeout to check for throttling/stop
                            )

                            if not done:
                                continue

                            # Process completed futures
                            for future in done:
                                file_path, file_type, file_name = futures[future]

                                try:
                                    result = future.result()

                                    # Put result in queue for batch processing
                                    results_queue.put(result)

                                    if result["success"]:
                                        successful_count += 1
                                        total_chunks += len(result["chunks"])
                                    else:
                                        failed_count += 1
                                        error_info = {
                                            "filename": file_name,
                                            "error": result["error"]
                                        }
                                        processing_errors.append(error_info)
                                        st.session_state.ingestion_errors.append(error_info)

                                except Exception as e:
                                    failed_count += 1
                                    error_info = {
                                        "filename": file_name,
                                        "error": f"Unhandled exception: {str(e)}"
                                    }
                                    processing_errors.append(error_info)
                                    st.session_state.ingestion_errors.append(error_info)

                                # Remove the completed future
                                del futures[future]

                                # Submit a new task if available
                                if next_file_index < len(all_files) and not st.session_state.stop_ingestion:
                                    new_future = executor.submit(process_file_for_parallel, all_files[next_file_index])
                                    futures[new_future] = all_files[next_file_index]
                                    next_file_index += 1

                                # Update progress
                                processed_count += 1
                                progress = processed_count / total_files
                                progress_bar.progress(progress)

                                # Calculate stats
                                elapsed = time.time() - start_time
                                effective_time = elapsed - total_throttle_time
                                files_per_second = processed_count / effective_time if effective_time > 0 else 0
                                estimated_total = effective_time / progress if progress > 0 else 0
                                remaining = estimated_total - effective_time

                                # Update status continuously
                                status_text.text(
                                    f"Processing: {processed_count}/{total_files} files | "
                                    f"Speed: {files_per_second:.2f} files/sec | "
                                    f"Time remaining: {remaining/60:.1f} minutes"
                                )

                                # Update metrics continuously
                                processed_metric.metric("Processed", f"{processed_count}/{total_files}")
                                successful_metric.metric("Successful", successful_count)
                                failed_metric.metric("Failed", failed_count)
                                chunks_metric.metric("Total Chunks", total_chunks)

                    # Signal that we're done adding to the collection
                    results_queue.put("DONE")

                    # Wait for the add thread to finish
                    add_thread.join()

                    # Complete the progress bar
                    progress_bar.progress(1.0)

                    # Clear the stop button
                    stop_button_placeholder.empty()
                    st.session_state.ingestion_running = False

                    # Show final summary
                    if st.session_state.stop_ingestion:
                        st.warning(
                            f"Ingestion stopped by user after {processed_count} of {total_files} files. "
                            f"Successfully processed {successful_count} files ({total_chunks} total chunks)."
                        )
                    else:
                        total_time = time.time() - start_time
                        st.success(
                            f"Ingestion complete in {total_time/60:.2f} minutes. "
                            f"Successfully processed {successful_count} of {total_files} files "
                            f"({total_chunks} total chunks)."
                            + (f" Throttled {throttle_count} times." if throttle_count > 0 else "")
                        )

                    # Display recent errors directly in this tab
                    if failed_count > 0:
                        st.warning(f"{failed_count} files failed to process. See the 'Errors' tab for details.")

                        with st.expander(f"Show recent errors ({min(5, len(processing_errors))} of {len(processing_errors)})"):
                            for error in processing_errors[:5]:
                                st.markdown(f"**File:** {error['filename']}")
                                st.code(error['error'], language="text")

            except Exception as e:
                st.error(f"Error accessing folder: {str(e)}")
                st.code(traceback.format_exc())
                st.info("Troubleshooting tips: \n"
                       "1. Make sure the path exists and is spelled correctly\n"
                       "2. Check that you have permission to access the folder\n"
                       "3. Try using an absolute path without special characters or spaces\n"
                       "4. On macOS/Linux, run 'ls -la \"your/path\"' in terminal to verify")
                st.session_state.ingestion_running = False

    # If ingestion is running, show the stop button
    elif st.session_state.ingestion_running:
        stop_button_placeholder.button("Stop Ingestion", on_click=stop_ingestion, key="stop_ingestion_button_persistent")

# ----------------------------
# Tab 4: Ingest Markdown Files
# ----------------------------

with tab_ingest_md:
    st.header("Ingest Markdown Documents")
    uploaded_md_files = st.file_uploader("Upload Markdown files", type="md", accept_multiple_files=True)

    if uploaded_md_files:
        collection = client.get_or_create_collection("tutorial_docs")

        # Create a progress bar for overall progress
        progress_bar = st.progress(0.0)

        # Create a status area for current file
        current_file_text = st.empty()

        # Process files with progress tracking
        total_files = len(uploaded_md_files)
        successful_files = 0
        failed_files = 0
        total_chunks = 0

        for i, file in enumerate(uploaded_md_files):
            # Update current file display
            current_file_text.text(f"Ingesting: {file.name} ({i+1}/{total_files})")

            # Update progress bar
            progress_bar.progress((i) / total_files)

            # Process the file
            result = process_file(file, collection, "markdown")

            if result["success"]:
                successful_files += 1
                total_chunks += result["chunks"]
            else:
                failed_files += 1
                st.session_state.ingestion_errors.append({
                    "filename": file.name,
                    "error": result["error"]
                })

            # Update progress bar
            progress_bar.progress((i+1) / total_files)

        # Clear the current file text
        current_file_text.empty()

        # Complete the progress bar
        progress_bar.progress(1.0)

        # Show final summary
        st.info(f"Ingestion complete. Successfully processed {successful_files} of {total_files} files ({total_chunks} total chunks).")

        if failed_files > 0:
            st.warning(f"{failed_files} files failed to process. See the 'Errors' tab for details.")

# ----------------------------
# Tab 5: Verify Document Count
# ----------------------------

with tab_verify:
    st.header("Verify Stored Document Count")

    try:
        all_collections = client.list_collections()
        st.subheader(f"Total collections: {len(all_collections)}")

        for coll_info in all_collections:
            name = coll_info.name
            col_key = f"delete_{name}"

            with st.container():
                col = client.get_or_create_collection(name)
                count = col.count()

                st.markdown(f"**Collection:** `{name}` — **Documents:** {count}")

                if st.button(f"Delete Collection '{name}'", key=col_key):
                    client.delete_collection(name)
                    st.success(f"Collection '{name}' deleted.")
                    st.experimental_rerun()

                # Load up to 100 documents for filtering and preview
                preview = col.get(include=["metadatas", "documents"], limit=100)
                all_metadatas = preview.get("metadatas", [])
                available_keys = {k for meta in all_metadatas if isinstance(meta, dict) for k in meta.keys()}

                # Optional filtering by metadata key/value
                if available_keys:
                    filter_key = st.selectbox(f"Filter by metadata key in '{name}'", sorted(available_keys), key=f"fk_{name}")
                    filter_value = st.text_input(f"Value to match for `{filter_key}`:", key=f"fv_{name}")

                    filtered_docs = []
                    filtered_ids = []
                    for i, meta in enumerate(all_metadatas):
                        if meta.get(filter_key) == filter_value:
                            filtered_docs.append(preview["documents"][i])
                            filtered_ids.append(preview["ids"][i])
                else:
                    filtered_docs = preview["documents"]
                    filtered_ids = preview["ids"]

                with st.expander(f"Show preview for collection '{name}'"):
                    for i in range(min(5, len(filtered_docs))):
                        st.markdown(f"**ID:** `{filtered_ids[i]}`")
                        st.code(filtered_docs[i])
                        if i < len(all_metadatas):
                            meta = all_metadatas[i]
                            if meta:
                                st.markdown("Metadata:")
                                st.json(meta)

    except Exception as e:
        st.error("Error retrieving collection info.")
        st.code(str(e))

# ----------------------------
# Tab 6: Ingestion Errors
# ----------------------------

with tab_errors:
    st.header("Ingestion Errors")

    if not st.session_state.ingestion_errors:
        st.info("No errors recorded during ingestion.")
    else:
        error_count = len(st.session_state.ingestion_errors)
        st.warning(f"{error_count} errors recorded during ingestion.")

        if st.button("Clear Error Log"):
            st.session_state.ingestion_errors = []
            st.experimental_rerun()

        # Group errors by type for a summary view
        error_types = {}
        for error in st.session_state.ingestion_errors:
            error_msg = error["error"].split('\n')[0]  # Get first line of error
            if error_msg in error_types:
                error_types[error_msg].append(error["filename"])
            else:
                error_types[error_msg] = [error["filename"]]

        # Display error summary
        st.subheader("Error Summary")
        for error_msg, filenames in error_types.items():
            with st.expander(f"{error_msg} ({len(filenames)} files)"):
                st.write("Affected files:")
                for filename in filenames:
                    st.write(f"- {filename}")

        # Option to view detailed errors
        if st.checkbox("Show detailed error logs"):
            for i, error in enumerate(st.session_state.ingestion_errors):
                with st.expander(f"Error details for: {error['filename']}"):
                    st.code(error['error'])

